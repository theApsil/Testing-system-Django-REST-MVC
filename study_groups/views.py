import os
from datetime import datetime

from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.core.files.storage import default_storage
from django.views.decorators.csrf import csrf_exempt
from rest_framework import status
import json
from docx import Document
from rest_framework.decorators import api_view, permission_classes
from rest_framework.parsers import MultiPartParser, FormParser
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework.views import APIView
from .models import Test, Question, Answer, UserAnswer, UserTestStatus, StudyGroup
from .serializers import TestSerializer, UserTestStatusSerializer, UserSerializer, QuestionSerializer
from rest_framework import generics
from .serializers import StudyGroupSerializer
from rest_framework.permissions import IsAuthenticated
import logging

logger = logging.getLogger(__name__)


class UserStudyGroupsView(generics.ListAPIView):
    serializer_class = StudyGroupSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return user.study_groups.all()


class TestDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, test_id):
        test = get_object_or_404(Test, id=test_id)
        serializer = TestSerializer(test)
        return Response(serializer.data)


class UserAnswerView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user = request.user
        data = request.data
        question = get_object_or_404(Question, id=data['question_id'])
        answer = get_object_or_404(Answer, id=data['answer_id'])
        UserAnswer.objects.update_or_create(user=user, question=question, defaults={'selected_answer': answer})
        return Response({"message": "Answer saved"}, status=status.HTTP_200_OK)


class UserResultsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        test_statuses = UserTestStatus.objects.filter(user=user, is_completed=True)
        results = []
        for test_status in test_statuses:
            user_answers = UserAnswer.objects.filter(user=user, question__test=test_status.test).values_list(
                'selected_answer_id', flat=True)
            result = {
                'test': TestSerializer(test_status.test).data,
                'score': test_status.score,
                'grade': test_status.grade,
                'user_answers': list(user_answers)
            }
            results.append(result)
        return Response(results)


class ActiveTestsView(generics.ListAPIView):
    serializer_class = TestSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        # Get all groups user is a member of
        groups = user.study_groups.all()
        # Get all tests that belong to these groups and user has not completed
        return Test.objects.filter(
            study_group__in=groups,
            userteststatus__is_completed=False,
            userteststatus__user=user
        )


class CompleteTestView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, test_id):
        user = request.user
        test = get_object_or_404(Test, id=test_id)
        user_answers = UserAnswer.objects.filter(user=user, question__test=test)
        correct_answers = 0
        for user_answer in user_answers:
            if user_answer.selected_answer.is_correct:
                correct_answers += 1
        score = correct_answers
        grade = 2
        total_questions = test.questions.count()
        if total_questions > 0:
            percentage = (correct_answers / total_questions) * 100
            if percentage >= 86:
                grade = 5
            elif percentage >= 70:
                grade = 4
            elif percentage >= 50:
                grade = 3

        UserTestStatus.objects.update_or_create(user=user, test=test,
                                                defaults={'is_completed': True, 'score': score, 'grade': grade})
        return Response({"message": "Test completed", "score": score, "grade": grade}, status=status.HTTP_200_OK)


class ProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        serializer = UserSerializer(user)
        return Response(serializer.data)


class TestsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        tests = Test.objects.all()
        serializer = TestSerializer(tests, many=True)
        return Response(serializer.data)


class TeacherGroupsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        if not user.is_staff:
            return Response({"error": "User is not a teacher"}, status=403)
        groups = StudyGroup.objects.filter(teacher_id=user.id)
        serializer = StudyGroupSerializer(groups, many=True)
        return Response(serializer.data)


class GroupStudentsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, group_id):
        try:
            group = StudyGroup.objects.get(study_group_id=group_id)
            students = group.members.filter(is_staff=False)
            serializer = UserSerializer(students, many=True)
            group_data = {
                "study_group_name": group.study_group_name,
                "teacher": {
                    "first_name": group.teacher.first_name,
                    "last_name": group.teacher.last_name
                },
                "members": serializer.data
            }
            return Response(group_data)
        except StudyGroup.DoesNotExist:
            return Response({"error": "Group not found"}, status=404)


class GroupTestsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, group_id):
        try:
            group = StudyGroup.objects.get(study_group_id=group_id)
            tests = group.tests.all()
            serializer = TestSerializer(tests, many=True)
            return Response(serializer.data)
        except StudyGroup.DoesNotExist:
            return Response({"error": "Group not found"}, status=404)


class StudentTestResultView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, student_id, test_id):
        try:
            test_status = UserTestStatus.objects.get(user_id=student_id, test_id=test_id)
            if not test_status.is_completed:
                return Response({"completed": False})

            user_answers = UserAnswer.objects.filter(user_id=student_id, question__test_id=test_id)
            questions = []
            for answer in user_answers:
                question_data = {
                    "text": answer.question.text,
                    "answer": answer.selected_answer.text,
                    "is_correct": answer.selected_answer.is_correct,
                }
                questions.append(question_data)

            result_data = {
                "completed": True,
                "score": test_status.score,
                "grade": test_status.grade,
                "questions": questions
            }
            logger.info(f"Result data: {result_data}")  # Debugging line
            return Response(result_data)
        except UserTestStatus.DoesNotExist:
            logger.error(f"Test status not found for user {student_id} and test {test_id}")
            return Response({"error": "Test status not found"}, status=404)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return Response({"error": str(e)}, status=500)


class GenerateReportView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, student_id, test_id, *args, **kwargs):
        try:
            test_status = UserTestStatus.objects.get(user_id=student_id, test_id=test_id)
            test = test_status.test
            user = test_status.user
        except UserTestStatus.DoesNotExist:
            return Response({"error": "UserTestStatus not found."}, status=status.HTTP_404_NOT_FOUND)

        # Путь к шаблону
        template_path = os.path.join('docs', 'analytics.docx')
        document = Document(template_path)

        # Замена переменных в документе
        self.replace_variable(document, 'DATE', datetime.now().strftime('%d.%m.%Y'))
        self.replace_variable(document, 'TEST', test.title)
        self.replace_variable(document, 'Фамилия', user.last_name)
        self.replace_variable(document, 'Имя', user.first_name)
        self.replace_variable(document, 'Балл', str(test_status.score))
        self.replace_variable(document, 'Оценка', str(test_status.grade))

        # Список ошибок
        errs = []

        user_answers = UserAnswer.objects.filter(user_id=student_id, question__test_id=test_id)
        for user_answer in user_answers:
            correct_answer = Answer.objects.filter(question=user_answer.question, is_correct=True).first()
            if user_answer.selected_answer != correct_answer:
                error = ''
                error += f"Вопрос: {user_answer.question.text} {chr(10)}"
                error += f"Ваш ответ: {user_answer.selected_answer.text} {chr(10)}"
                error += f"Правильный ответ: {correct_answer.text}"
                errs.append(error)

        if errs:
            text = ("Список ошибок:\n" +
                    "\n========================================================\n".join(errs))
        else:
            text = "Тест пройден без ошибок."

        self.replace_variable(document, 'Список ошибок', text)

        # Создание временного файла
        temp_file = 'temp_report.docx'
        document.save(temp_file)

        # Возвращение документа в качестве ответа
        with open(temp_file, 'rb') as f:
            response = HttpResponse(f.read(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="analytics_report.docx"'
            return response

    def replace_variable(self, document, variable, value):
        for paragraph in document.paragraphs:
            if variable in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if variable in item.text:
                        item.text = item.text.replace(variable, value)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_variable(cell, variable, value)


class CreateQuestionView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        text = request.data.get('text')
        image = request.FILES.get('image')
        answers_data = json.loads(request.data.get('answers'))

        question = Question.objects.create(text=text, image=image)

        for answer_data in answers_data:
            Answer.objects.create(
                question=question,
                text=answer_data['text'],
                is_correct=answer_data['is_correct']
            )

        return Response({'message': 'Задание успешно создано'}, status=status.HTTP_201_CREATED)


class AllQuestionsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        questions = Question.objects.all()
        serializer = QuestionSerializer(questions, many=True)
        return Response(serializer.data)


class CreateTestView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        title = request.data.get('title')
        description = request.data.get('description')
        study_group_id = request.data.get('study_group')
        question_ids = request.data.get('questions', [])

        try:
            study_group = StudyGroup.objects.get(pk=study_group_id)
        except StudyGroup.DoesNotExist:
            return Response({'error': 'Invalid study_group.'}, status=status.HTTP_400_BAD_REQUEST)

        if not title or not description:
            return Response({'error': 'Title and description are required.'}, status=status.HTTP_400_BAD_REQUEST)

        # Создание теста
        test = Test.objects.create(
            title=title,
            description=description,
            cover_image='/static/img/img_4.png',
            study_group_id=study_group_id
        )

        questions = Question.objects.filter(id__in=question_ids)
        test.questions.set(questions)
        test.save()

        for member in study_group.members.all():
            UserTestStatus.objects.create(
                test=test,
                user=member,
                is_completed=False,
                grade=0,
                score=0
            )

        serializer = TestSerializer(test)
        return Response(serializer.data, status=status.HTTP_201_CREATED)